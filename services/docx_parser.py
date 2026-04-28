"""DOCX 파일을 서식 메타데이터와 함께 파싱하는 모듈.

python-docx로 각 paragraph의 스타일, 서식(굵기/기울임/밑줄/취소선/대문자),
들여쓰기 레벨, 목록 레벨, 표 데이터를 추출합니다.
"""

import re as _re
from dataclasses import dataclass, field
from typing import List, Optional, Union
from docx import Document

# Word XML 네임스페이스
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
NSMAP = {'w': WORD_NS}


@dataclass
class DocxRun:
    """paragraph 내 개별 텍스트 런"""
    text: str
    is_bold: bool = False
    is_italic: bool = False
    is_underline: bool = False
    is_strike: bool = False


@dataclass
class DocxParagraph:
    """서식 메타데이터를 포함한 paragraph"""
    text: str
    style_name: str = "Normal"
    is_bold: bool = False          # 전체 paragraph가 굵은지
    is_italic: bool = False
    is_underline: bool = False
    is_strikethrough: bool = False
    is_all_caps: bool = False
    indent_level: int = 0          # left_indent 기반 (0, 1, 2...)
    list_level: Optional[int] = None   # 목록 레벨 (None이면 목록 아님)
    list_num_id: Optional[int] = None  # 같은 목록 그룹 ID
    is_numbered_list: bool = False     # True면 번호 목록, False면 글머리 기호
    runs: List[DocxRun] = field(default_factory=list)


@dataclass
class DocxCell:
    """표의 개별 셀 — 병합 정보 포함."""
    text: str
    is_merged_continuation: bool = False  # True = 병합으로 인한 복제 셀
    row_span: int = 1
    col_span: int = 1


@dataclass
class DocxTable:
    """DOCX 테이블"""
    rows: List[List[str]] = field(default_factory=list)       # 하위 호환용 유지
    header_row: List[str] = field(default_factory=list)
    row_count: int = 0
    col_count: int = 0
    table_type: str = "unknown"
    has_merged_cells: bool = False
    raw_cells: List[List['DocxCell']] = field(default_factory=list)

    @property
    def rows_text(self) -> List[List[str]]:
        """continuation 셀 제거한 클린 행 텍스트."""
        if not self.raw_cells:
            return self.rows
        return [
            [c.text for c in row if not c.is_merged_continuation]
            for row in self.raw_cells
        ]


ContentItem = Union[DocxParagraph, DocxTable]


@dataclass
class DocxSection:
    """문서의 논리적 섹션 (Heading 단위로 분리)"""
    heading: Optional[str] = None
    content: List[ContentItem] = field(default_factory=list)  # 단락+표 원본 순서 보존

    @property
    def paragraphs(self) -> List[DocxParagraph]:
        return [c for c in self.content if isinstance(c, DocxParagraph)]

    @property
    def tables(self) -> List[DocxTable]:
        return [c for c in self.content if isinstance(c, DocxTable)]


def _extract_textbox_text(paragraph) -> str:
    """paragraph 내 텍스트 박스(wps:txbx)의 텍스트를 추출.

    Returns:
        추출된 텍스트 또는 빈 문자열 (텍스트 박스 없으면)
    """
    try:
        for drawing in paragraph._element.findall(
            f'.//{{{WORD_NS}}}drawing'
        ):
            for txbx_content in drawing.findall(
                f'.//{{{WPS_NS}}}txbx/{{{WORD_NS}}}txbxContent'
            ):
                texts = []
                for p in txbx_content.findall(f'{{{WORD_NS}}}p'):
                    t = ''.join(
                        r.text or '' for r in p.findall(f'.//{{{WORD_NS}}}t')
                    ).strip()
                    if t:
                        texts.append(t)
                if texts:
                    return ' / '.join(texts)
    except Exception:
        pass
    return ''


def _get_list_info(paragraph) -> tuple:
    """paragraph의 XML에서 목록 레벨과 numId를 추출.

    Returns:
        (list_level, num_id, is_numbered)
    """
    pPr = paragraph._element.find(f'{{{WORD_NS}}}pPr')
    if pPr is None:
        return None, None, False

    numPr = pPr.find(f'{{{WORD_NS}}}numPr')
    if numPr is None:
        return None, None, False

    ilvl_elem = numPr.find(f'{{{WORD_NS}}}ilvl')
    numId_elem = numPr.find(f'{{{WORD_NS}}}numId')

    list_level = int(ilvl_elem.get(f'{{{WORD_NS}}}val', '0')) if ilvl_elem is not None else 0
    num_id = int(numId_elem.get(f'{{{WORD_NS}}}val', '0')) if numId_elem is not None else None

    # numId == 0은 목록 아님
    if num_id == 0:
        return None, None, False

    # 번호 매기기 형식 확인 시도
    is_numbered = _check_numbering_format(paragraph, num_id, list_level)
    return list_level, num_id, is_numbered


def _check_numbering_format(paragraph, num_id, list_level) -> bool:
    """numbering.xml에서 번호 매기기 형식 확인.

    Returns: True면 번호 목록, False면 글머리 기호
    """
    try:
        numbering_part = paragraph.part.numbering_part
        if numbering_part is None:
            return False

        numbering_xml = numbering_part._element

        # w:num 요소 찾기
        for num_elem in numbering_xml.findall(f'{{{WORD_NS}}}num'):
            if int(num_elem.get(f'{{{WORD_NS}}}numId', '0')) == num_id:
                abstract_num_id_elem = num_elem.find(f'{{{WORD_NS}}}abstractNumId')
                if abstract_num_id_elem is None:
                    continue
                abstract_num_id = int(abstract_num_id_elem.get(f'{{{WORD_NS}}}val', '0'))

                # w:abstractNum에서 numFmt 확인
                for abstract_num in numbering_xml.findall(f'{{{WORD_NS}}}abstractNum'):
                    if int(abstract_num.get(f'{{{WORD_NS}}}abstractNumId', '0')) == abstract_num_id:
                        for lvl in abstract_num.findall(f'{{{WORD_NS}}}lvl'):
                            if int(lvl.get(f'{{{WORD_NS}}}ilvl', '0')) == list_level:
                                num_fmt = lvl.find(f'{{{WORD_NS}}}numFmt')
                                if num_fmt is not None:
                                    fmt_val = num_fmt.get(f'{{{WORD_NS}}}val', '')
                                    # bullet은 글머리 기호, 나머지(decimal, lowerLetter 등)는 번호
                                    return fmt_val != 'bullet'
                        break
                break
    except Exception:
        pass
    return False


def _get_indent_level(paragraph) -> int:
    """paragraph의 들여쓰기 레벨 계산 (EMU → 레벨)"""
    pf = paragraph.paragraph_format
    left_indent = pf.left_indent
    if left_indent is None:
        return 0
    # 1 레벨 ≈ 360000 EMU (약 0.25인치)
    return max(0, int(left_indent / 360000))


def _parse_paragraph(paragraph) -> Optional[DocxParagraph]:
    """python-docx paragraph를 DocxParagraph로 변환"""
    text = paragraph.text.strip()

    # 빈 paragraph도 텍스트 박스가 있을 수 있으므로 먼저 확인
    if not text:
        textbox_text = _extract_textbox_text(paragraph)
        if textbox_text:
            text = f"[TEXTBOX: {textbox_text}]"
        else:
            return None

    # 스타일 이름
    style_name = paragraph.style.name if paragraph.style else "Normal"

    # 런 분석
    runs = []
    all_bold = True
    all_italic = True
    all_underline = True
    all_strike = True
    all_caps = True
    has_runs = False

    for run in paragraph.runs:
        if not run.text.strip():
            continue
        has_runs = True

        run_bold = bool(run.font.bold)
        run_italic = bool(run.font.italic)
        run_underline = bool(run.font.underline)
        run_strike = bool(run.font.strike)

        if not run_bold:
            all_bold = False
        if not run_italic:
            all_italic = False
        if not run_underline:
            all_underline = False
        if not run_strike:
            all_strike = False
        if not run.font.all_caps:
            all_caps = False

        runs.append(DocxRun(
            text=run.text,
            is_bold=run_bold,
            is_italic=run_italic,
            is_underline=run_underline,
            is_strike=run_strike,
        ))

    if not has_runs:
        all_bold = False
        all_italic = False
        all_underline = False
        all_strike = False
        all_caps = False

    # 취소선 paragraph는 건너뛰기
    if all_strike and has_runs:
        return None

    # 목록 정보
    list_level, list_num_id, is_numbered = _get_list_info(paragraph)

    # 들여쓰기 레벨
    indent_level = _get_indent_level(paragraph)

    # 텍스트 박스 추출
    textbox_text = _extract_textbox_text(paragraph)
    if textbox_text:
        if text:
            text = f"{text} [TEXTBOX: {textbox_text}]"
        else:
            text = f"[TEXTBOX: {textbox_text}]"

    return DocxParagraph(
        text=text,
        style_name=style_name,
        is_bold=all_bold,
        is_italic=all_italic,
        is_underline=all_underline,
        is_strikethrough=all_strike,
        is_all_caps=all_caps,
        indent_level=indent_level,
        list_level=list_level,
        list_num_id=list_num_id,
        is_numbered_list=is_numbered,
        runs=runs,
    )


# 코딩 참고 표 헤더 키워드
_CODING_REF_HEADERS = frozenset({
    'variable', 'var', 'code', 'label', 'value', 'name', 'description',
    '변수', '코드', '레이블', '값', '이름', '설명', '항목',
})

# 문항 번호 패턴 (multi_question 감지용)
_QN_IN_CELL_RE = _re.compile(
    r'^[A-Za-z]+\d+[a-z]?[.)\s]',
)

_TYPE_HINT_RE = _re.compile(
    r'^\s*\[(?:SA|S|MA|M|OE|OPEN|NUMERIC|SCALE|RANK|TOP\s*\d+)[^\]]*\]\s*$',
    _re.IGNORECASE,
)

_CODE_CELL_RE = _re.compile(
    r'^\s*(?:-?\d+[.)]?|[A-Za-z]|[A-Za-z]\d{1,2})\s*$'
)

_ROUTING_CELL_RE = _re.compile(
    r'(?:terminate|screen\s*out|go\s*to|skip\s*to|quota|종료|탈락|이동|건너뜀)',
    _re.IGNORECASE,
)

_CODING_REF_STRONG_HEADERS = frozenset({
    'variable', 'var', 'name', 'description', 'quota', 'criteria',
    '변수', '이름', '설명', '항목', '쿼터', '조건',
})


def _clean_table_rows(rows: List[List[str]]) -> List[List[str]]:
    """Trim cells and drop fully empty rows for table classification."""
    cleaned = []
    for row in rows:
        cells = [str(c).strip() for c in row]
        if any(cells):
            cleaned.append(cells)
    return cleaned


def _is_code_cell(text: str) -> bool:
    """Return True for compact option/code cells such as 1, 99, A, B2."""
    if not text:
        return False
    value = text.strip()
    if _TYPE_HINT_RE.match(value):
        return False
    return bool(_CODE_CELL_RE.match(value))


def _looks_like_option_label(text: str) -> bool:
    """Return True for the label side of a response option row."""
    if not text:
        return False
    value = text.strip()
    if _TYPE_HINT_RE.match(value):
        return False
    if _is_code_cell(value):
        return False
    if _ROUTING_CELL_RE.fullmatch(value):
        return False
    return True


def _is_option_header(row: List[str]) -> bool:
    """Detect non-data header rows in option tables."""
    cells = [c.strip().lower() for c in row if c.strip()]
    if not cells:
        return True
    if all(_TYPE_HINT_RE.match(c) for c in row if c.strip()):
        return True
    header_terms = {'code', 'label', 'value', 'option', 'answer',
                    '코드', '레이블', '값', '보기', '응답'}
    return bool(cells) and all(c in header_terms for c in cells)


def _looks_like_code_label_table(rows: List[List[str]], col_count: int) -> bool:
    """Detect answer option tables in both code|label and label|code layouts.

    Many Ipsos questionnaires use label|code or label|code|routing columns,
    not only the classic code|label layout. This classifier intentionally
    accepts either orientation but rejects wide scale rows with many numeric
    scale columns.
    """
    if col_count < 2 or col_count > 4:
        return False

    data_rows = _clean_table_rows(rows)
    if data_rows and _is_option_header(data_rows[0]):
        data_rows = data_rows[1:]
    if not data_rows:
        return False

    matched = 0
    considered = 0
    for row in data_rows:
        nonempty = [(idx, cell.strip()) for idx, cell in enumerate(row) if cell.strip()]
        if len(nonempty) < 2:
            continue

        code_positions = [idx for idx, cell in nonempty if _is_code_cell(cell)]
        # Wide numeric rows are usually scales, not answer-option rows.
        if len(code_positions) > 2:
            continue

        considered += 1
        row_matched = False
        for code_idx in code_positions:
            before = [cell for idx, cell in nonempty if idx < code_idx]
            after = [cell for idx, cell in nonempty if idx > code_idx]
            label_candidates = before + after
            if any(_looks_like_option_label(cell) for cell in label_candidates):
                row_matched = True
                break
        if row_matched:
            matched += 1

    if considered == 0:
        return False
    return matched >= 2 and matched / considered >= 0.6


def _classify_table(rows: List[List[str]]) -> str:
    """표의 내용을 분석하여 타입을 분류한다.

    Returns:
        'section_header' | 'coding_reference' | 'multi_question' |
        'code_label' | 'grid' | 'matrix' | 'generic' | 'unknown'
    """
    if not rows:
        return "unknown"

    row_count = len(rows)
    col_count = max(len(r) for r in rows) if rows else 0

    # ── section_header: 1×1 표, 짧은 텍스트, 지시문 패턴 없음 ──
    if row_count == 1 and col_count == 1:
        text = rows[0][0].strip()
        if not text:
            return "unknown"
        # 100자 이상이면 지시문/설명일 가능성 → generic으로 보존
        if len(text) > 100:
            return "generic"
        # 지시문/필터 패턴이 포함되어 있으면 generic으로 보존
        text_upper = text.upper()
        if any(kw in text_upper for kw in (
            '[PN:', 'ASK IF', 'ASK ALL', 'SHOW CARD', 'ROTATE',
            '모두에게', '응답자', '제시',
        )):
            return "generic"
        return "section_header"

    # ── multi_question: 첫 번째 열에 여러 문항 번호 패턴 ──
    if col_count >= 2:
        first_col_values = [rows[i][0].strip() for i in range(row_count)
                            if rows[i] and rows[i][0].strip()]
        qn_count = sum(1 for v in first_col_values if _QN_IN_CELL_RE.match(v))
        if first_col_values and qn_count / len(first_col_values) >= 0.5 and qn_count >= 2:
            return "multi_question"

    # ── coding_reference: 변수/코드북/쿼터 등 명확한 메타데이터 표 ──
    if rows and col_count >= 2:
        header_cells = [c.strip().lower() for c in rows[0] if c.strip()]
        strong_matches = sum(1 for c in header_cells if c in _CODING_REF_STRONG_HEADERS)
        if header_cells and strong_matches > 0:
            return "coding_reference"

    # ── code_label: 보기 표 (코드+라벨 구조) ──
    # 2~4열, code|label / label|code / label|code|routing 모두 허용
    if row_count >= 2 and _looks_like_code_label_table(rows, col_count):
        return "code_label"

    # ── coding_reference: 3열 이상에서 Code/Label/Value류 헤더가 지배적인 메타데이터 표 ──
    if rows and col_count >= 2:
        header_cells = [c.strip().lower() for c in rows[0] if c.strip()]
        header_matches = sum(1 for c in header_cells if c in _CODING_REF_HEADERS)
        if header_cells and col_count >= 3 and header_matches / len(header_cells) >= 0.6:
            return "coding_reference"

    # ── code_label 폴백: 1행 가로 배치 보기 (| 1.남성 | 2.여성 | 3.기타 |) ──
    if row_count == 1 and col_count >= 3:
        cells = [c.strip() for c in rows[0] if c.strip()]
        coded_cells = sum(1 for c in cells if _re.match(r'^\d+[.)\s]', c))
        if cells and coded_cells / len(cells) >= 0.6:
            return "code_label"

    # ── grid: 첫 행이 숫자 척도, 나머지 행이 항목 ──
    if row_count >= 3 and col_count >= 3:
        header = [c.strip() for c in rows[0] if c.strip()]
        scale_headers = [c for c in header[1:] if c] if len(header) > 1 else header
        numeric_headers = sum(1 for c in scale_headers if c.isdigit() or
                              c in ('○', '●', '□', '■'))
        if scale_headers and numeric_headers / len(scale_headers) >= 0.6:
            return "grid"

    # ── matrix: 3행 이상, 3열 이상, 첫 열이 텍스트 항목 ──
    if row_count >= 3 and col_count >= 3:
        first_col = [rows[i][0].strip() for i in range(1, row_count) if rows[i]]
        non_numeric = sum(1 for v in first_col if v and not v.lstrip('-').isdigit())
        if first_col and non_numeric / len(first_col) >= 0.7:
            return "matrix"

    return "generic"


def _is_cell_strikethrough(cell) -> bool:
    """셀의 텍스트 런이 모두 취소선(strikethrough)인지 확인.

    취소선 셀 = 삭제된 보기. 빈 셀이면 False 반환.
    셀 내 여러 단락이 있으면 첫 번째 비어있지 않은 단락 기준으로 판단.
    """
    for para in cell.paragraphs:
        runs_with_text = [r for r in para.runs if r.text.strip()]
        if not runs_with_text:
            continue
        return all(bool(r.font.strike) for r in runs_with_text)
    return False


def _parse_table(table) -> DocxTable:
    """python-docx Table을 DocxTable로 변환.

    병합 셀은 seen_tc_ids로 중복 제거하고,
    취소선 셀은 삭제된 내용으로 간주하여 빈 문자열로 처리한다.
    """
    raw_cells_data: List[List[DocxCell]] = []
    rows_data: List[List[str]] = []
    has_merged = False

    for row in table.rows:
        seen_tc_ids: set = set()
        row_cells: List[DocxCell] = []
        row_text: List[str] = []

        for cell in row.cells:
            tc_id = id(cell._tc)
            is_dup = tc_id in seen_tc_ids
            seen_tc_ids.add(tc_id)

            if is_dup:
                has_merged = True
                row_cells.append(DocxCell(text=cell.text.strip(),
                                          is_merged_continuation=True))
            else:
                # 취소선 셀은 빈 문자열로 치환
                text = "" if _is_cell_strikethrough(cell) else cell.text.strip()
                row_cells.append(DocxCell(text=text,
                                          is_merged_continuation=False))
                row_text.append(text)

        raw_cells_data.append(row_cells)
        rows_data.append(row_text)

    table_type = _classify_table(rows_data)

    return DocxTable(
        rows=rows_data,
        header_row=rows_data[0] if rows_data else [],
        row_count=len(rows_data),
        col_count=len(rows_data[0]) if rows_data else 0,
        table_type=table_type,
        has_merged_cells=has_merged,
        raw_cells=raw_cells_data,
    )


def parse_docx(file) -> List[DocxSection]:
    """DOCX 파일을 서식 정보와 함께 섹션 단위로 파싱.

    Args:
        file: 업로드된 파일 객체 또는 파일 경로

    Returns:
        DocxSection 리스트 (Heading 단위로 분리)
    """
    doc = Document(file)
    sections = []
    current_section = DocxSection()

    # 문서 body의 모든 요소를 순서대로 처리
    # doc.element.body의 자식 요소를 순회하여 paragraph와 table 순서를 보존
    body = doc.element.body

    # paragraph와 table을 인덱싱하기 위한 매핑
    para_index = 0
    table_index = 0

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # Paragraph
            if para_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[para_index]
                para_index += 1

                parsed = _parse_paragraph(paragraph)
                if parsed is None:
                    continue

                # Heading이면 새 섹션 시작
                if parsed.style_name and 'Heading' in parsed.style_name:
                    if current_section.content or current_section.heading:
                        sections.append(current_section)
                    current_section = DocxSection(heading=parsed.text)
                else:
                    current_section.content.append(parsed)

        elif tag == 'tbl':
            # Table
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                table_index += 1

                parsed_table = _parse_table(table)
                if parsed_table.row_count == 0:
                    continue

                if parsed_table.table_type == "section_header":
                    heading_text = parsed_table.rows[0][0].strip()
                    if heading_text:
                        if current_section.content or current_section.heading:
                            sections.append(current_section)
                        current_section = DocxSection(heading=heading_text)
                else:
                    current_section.content.append(parsed_table)

    # 마지막 섹션 추가
    if current_section.content or current_section.heading:
        sections.append(current_section)

    return sections
